# app/chunks_creating.py
import uuid 
import re
from docx import Document as DocxDocument
from langchain.schema import Document
from app import photo_indexes
import os
from .preprocessing.context_annotator import ContextAnnotator
from langchain.text_splitter import RecursiveCharacterTextSplitter
from rank_bm25 import BM25Okapi
from typing import List
from tqdm import tqdm
# Чтение текста из файла .docx
def load_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Путь к твоему файлу .docx
file_path = "documentation_preprocessed.docx"

# Загружаем текст из файла
text_content = load_text_from_docx(file_path)

# titles из твоего словаря заголовков, включая Аннотацию и КОНТАКТНУЮ ИНФОРМАЦИЮ
titles = photo_indexes.titles
titles["Аннотация"] = "Аннотация"
titles["КОНТАКТНАЯ ИНФОРМАЦИЯ"] = "КОНТАКТНАЯ ИНФОРМАЦИЯ"

# Регулярное выражение для поиска заголовков, включая их название (добавлены "Аннотация" и "КОНТАКТНАЯ ИНФОРМАЦИЯ")
heading_pattern = re.compile(r"(\d+(\.\d+)*|Аннотация|КОНТАКТНАЯ ИНФОРМАЦИЯ)")

# Функция для разделения текста на части по заголовкам (включая название)
def split_text_by_headings(text: str, titles: dict):
    documents = []
    matches = list(heading_pattern.finditer(text))
    
    for i in range(len(matches)):
        start = matches[i].start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        
        chapter = matches[i].group(1).strip()
        chapter_title = titles.get(chapter, "Неизвестный заголовок")
        
        content = text[start:end].strip()
        image_paths = get_files_by_number(chapter)
        
        doc = Document(
            page_content=content[:1000],
            metadata={
                "id": str(uuid.uuid4()),  # Добавляем уникальный ID
                "content": content,
                "chapter": chapter,
                "title": chapter_title,
                "paths": image_paths
            }
        )
        documents.append(doc)
    
    return documents

# Получение файлов по номеру главы (реализуй эту функцию)
def get_files_by_number(number):
    # Здесь функция get_files_by_number уже должна быть реализована
    # для получения путей к фоткам, которые соответствуют главе
    return [f for f in os.listdir("images") if f.startswith(str(number)+'_')]

# Применяем функцию для разделения текста
async def create_annotated_documents() -> List[Document]:
    # Создаем базовые документы как раньше
    documents = split_text_by_headings(text_content, titles)
    
    # Создаем аннотатор
    annotator = ContextAnnotator()
    
    # Добавляем контекст к каждому документу
    annotated_documents = []
    for doc in tqdm(documents):
        # annotated_text = await annotator.annotate_chunk(
        #     doc.page_content,
        #     doc.metadata
        # )
        annotated_text = "Текст про " + doc.page_content 
        # Store both original content and annotated text
        original_content = doc.page_content
        doc.page_content = annotated_text  # annotated for vectorstore
        annotated_documents.append(doc)
        
        # Create a separate Document for BM25 indexing with original content
        bm25_doc = Document(
            page_content=original_content,
            metadata=doc.metadata  # share metadata
        )
        annotated_documents.append(bm25_doc)
        # annotated_documents = []
    return annotated_documents

def create_bm25_index(documents: List[Document]) -> tuple:
    # Separate original documents for BM25
    original_docs = [doc for doc in documents if doc.page_content == doc.metadata['content']]
    
    # Создаем BM25 индекс
    tokenized_docs = [doc.page_content.split() for doc in original_docs]
    bm25 = BM25Okapi(tokenized_docs)
    
    # Добавляем BM25 скоры в метаданные документов
    for i, doc in enumerate(original_docs):
        doc.metadata['bm25_scores'] = bm25.get_scores(doc.page_content.split())
    
    return original_docs, bm25
